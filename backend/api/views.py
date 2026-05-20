import os
import uuid
import requests
from django.conf import settings
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. إعداد مسار Tesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# 2. إعداد المفتاح
GEMINI_API_KEY = 'AIzaSyC_-_c7VOQ79KOVLMPDefoY5a272wdNn6k'

def get_available_model():
    """هذه الدالة تستخرج النموذج الصحيح والمسموح به لمفتاحك تلقائياً"""
    url = f"https://generativelanguage.googleapis.com/v1beta/models?key={GEMINI_API_KEY}"
    response = requests.get(url)
    if response.status_code == 200:
        models = response.json().get('models', [])
        # نبحث عن نموذج يدعم توليد النصوص (generateContent)
        for model in models:
            if 'generateContent' in model.get('supportedGenerationMethods', []):
                # نفضل نماذج flash لأنها أسرع وأفضل للنصوص
                if 'flash' in model['name']:
                    return model['name']
        # إذا لم يجد flash، يرجع أول نموذج متاح
        for model in models:
             if 'generateContent' in model.get('supportedGenerationMethods', []):
                 return model['name']
    return "models/gemini-1.5-flash" # نموذج احتياطي أخير

class PDFToWordView(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        pdf_file = request.FILES.get('pdf')
        if not pdf_file:
            return Response({'error': 'لم يتم اختيار ملف'}, status=400)

        unique_id = uuid.uuid4().hex
        pdf_path = os.path.join(settings.MEDIA_ROOT, f"temp_{unique_id}.pdf")
        docx_path = os.path.join(settings.MEDIA_ROOT, f"converted_{unique_id}.docx")

        with open(pdf_path, 'wb+') as f:
            for chunk in pdf_file.chunks():
                f.write(chunk)

        try:
            # 1. قراءة الـ PDF بالـ OCR
            pages = convert_from_path(pdf_path, 300, poppler_path=r'C:\poppler-26.02.0\Library\bin')
            extracted_text = ""
            
            for page in pages:
                text = pytesseract.image_to_string(page, lang='ara')
                extracted_text += text + "\n"

            # 2. تحديد النموذج الديناميكي
            model_name = get_available_model()
            
            prompt = f"""
            أنت سكرتير تنفيذي محترف. النص التالي تم استخراجه من تقرير أمني عربي عبر OCR ويحتوي على أخطاء وتنسيق عشوائي.
            
            المطلوب:
            1. صحح الأخطاء الإملائية.
            2. رتب التقرير بشكل هرمي واضح.
            3. استخدم الترقيم (1.  أو 2. ) للعناوين، واستخدم الشرطة (-) للبنود الفرعية.
            4. تأكد أن النص نظيف تماماً.
            5. أعد النص فقط بدون أي تعليقات أو علامات ماركداون (** أو #).

            النص:
            {extracted_text}
            """

            # 3. إرسال الطلب باستخدام النموذج المكتشف
            url = f"https://generativelanguage.googleapis.com/v1beta/{model_name}:generateContent?key={GEMINI_API_KEY}"
            headers = {'Content-Type': 'application/json'}
            data = {
                "contents": [{"parts": [{"text": prompt}]}]
            }

            response = requests.post(url, headers=headers, json=data)
            response_data = response.json()

            if response.status_code != 200:
                error_msg = response_data.get('error', {}).get('message', 'Unknown Error')
                raise Exception(f"Google API Error: {error_msg}")

            clean_formatted_text = response_data['candidates'][0]['content']['parts'][0]['text']

            # 4. كتابة ملف الوورد مع دعم الاتجاه العربي (RTL)
            doc = Document()
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            lines = clean_formatted_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                clean_line = line.replace('**', '').replace('#', '').strip()
                clean_line = '\u200F' + clean_line

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                pPr = p._element.get_or_add_pPr()
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                pPr.append(bidi)
                
                run = p.add_run(clean_line)
                
                rPr = run._element.get_or_add_rPr()
                rtl = OxmlElement('w:rtl')
                rtl.set(qn('w:val'), '1')
                rPr.append(rtl)
                
                lang = OxmlElement('w:lang')
                lang.set(qn('w:bidi'), 'ar-SA') 
                rPr.append(lang)
                
                run.font.name = 'Arial' 

                if clean_line and len(clean_line) > 1 and (clean_line[1].isdigit() or "وزارة" in clean_line or "قيادة" in clean_line or "الملخص" in clean_line):
                    run.bold = True
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)

            doc.save(docx_path)
            file_url = request.build_absolute_uri(f"{settings.MEDIA_URL}converted_{unique_id}.docx")
            return Response({'docx_url': file_url})

        except Exception as e:
            print(f"Error Output: {str(e)}")
            return Response({'error': f'فشل التحويل: {str(e)}'}, status=500)
        
        finally:
            if os.path.exists(pdf_path):
                try:
                    os.remove(pdf_path)
                except:
                    pass